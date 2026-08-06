import fitz
import pytest
from docx import Document as DocxWriter

from src.document_loader import (
	TxtDocumentLoader,
	PdfDocumentLoader,
	DocxDocumentLoader,
	HtmlDocumentLoader,
)
from src.document_factory import DocumentLoaderFactory
from src.document_exceptions import DocumentParsingError, UnsupportedFormatError


# DocumentLoader base behaviour
def test_loader_raises_filenotfound_for_missing_file(tmp_path):
	missing_path = tmp_path / "does_not_exist.txt"

	with pytest.raises(FileNotFoundError):
		TxtDocumentLoader(missing_path)

# TxtDocumentLoader
def test_txt_loader_reads_utf8_content(tmp_path):
	file_path = tmp_path / "note.txt"
	file_path.write_text("hello world", encoding="utf-8")
	
	documents = TxtDocumentLoader(file_path).load()
	
	assert len(documents) == 1
	assert documents[0].content == "hello world"
	assert documents[0].document_type == "txt"
	assert documents[0].metadata.page_number == 1
	assert documents[0].metadata.filename == "note.txt"
	
def test_txt_loader_autodetects_non_utf8_encoding(tmp_path):
	file_path = tmp_path / "latin.txt"
	file_path.write_text("Lorem ipsum", encoding="latin-1")
	documents = TxtDocumentLoader(file_path).load()
	# test only if the file is read without exceptions and produces non-empty content.
	assert documents[0].content.strip() != ""


# PdfDocumentLoader
def _make_pdf(path, pages_text: list[str]) -> None:
	doc = fitz.open()
	for text in pages_text:
		page = doc.new_page()
		if text:
			page.insert_text((72, 72), text)
	doc.save(str(path))
	doc.close()

def test_pdf_loader_extracts_text_per_page(tmp_path):
	pdf_path = tmp_path / "doc.pdf"
	_make_pdf(pdf_path, ["page one", "page two"])
	documents = PdfDocumentLoader(pdf_path).load()

	assert len(documents) == 2
	assert documents[0].metadata.page_number == 1
	assert documents[1].metadata.page_number == 2
	assert "page one" in documents[0].content
	assert "page two" in documents[1].content

def test_pdf_loader_skips_empty_pages(tmp_path):
	pdf_path = tmp_path / "doc.pdf"
	_make_pdf(pdf_path, ["Lorem ipsum", ""])
	documents = PdfDocumentLoader(pdf_path).load()

	# the blank page must not produce a LoadedDocument
	assert len(documents) == 1
	assert documents[0].metadata.page_number == 1

def test_pdf_loader_raises_parsing_error_on_corrupted_file(tmp_path):
	fake_pdf = tmp_path / "corrupted.pdf"
	fake_pdf.write_bytes(b"This is not a real PDF")
	with pytest.raises(DocumentParsingError):
		PdfDocumentLoader(fake_pdf).load()


# DocxDocumentLoader
def test_docx_loader_extracts_paragraphs_and_tables_in_order(tmp_path):
	docx_path = tmp_path / "doc.docx"
	doc = DocxWriter()
	doc.add_paragraph("Title")
	table = doc.add_table(rows=1, cols=2)
	table.cell(0, 0).text = "Name"
	table.cell(0, 1).text = "Value"
	doc.add_paragraph("Text after the table")
	doc.save(str(docx_path))

	documents = DocxDocumentLoader(docx_path).load()

	assert len(documents) == 1
	content = documents[0].content
	# the order in the extracted text must reflect the actual order in the document
	assert content.index("Title") < content.index("Name | Value")
	assert content.index("Name | Value") < content.index("Text after the table")


def test_docx_loader_raises_parsing_error_on_corrupted_file(tmp_path):
	fake_docx = tmp_path / "corrupted.docx"
	fake_docx.write_bytes(b"not a real docx")
	with pytest.raises(DocumentParsingError):
		DocxDocumentLoader(fake_docx).load()


# HtmlDocumentLoader
def test_html_loader_strips_tags_and_keeps_text(tmp_path):
	html_path = tmp_path / "page.html"
	html_path.write_text(
		"<html><body><h1>Title</h1><p>A paragraph.</p></body></html>",
		encoding="utf-8",
	)
	documents = HtmlDocumentLoader(html_path).load()

	assert "Title" in documents[0].content
	assert "A paragraph." in documents[0].content
	assert "<h1>" not in documents[0].content


# DocumentLoaderFactory
@pytest.mark.parametrize(
	"extension, expected_loader_type",
	[
		(".txt", TxtDocumentLoader),
		(".md", TxtDocumentLoader),
		(".pdf", PdfDocumentLoader),
		(".docx", DocxDocumentLoader),
		(".html", HtmlDocumentLoader),
		(".htm", HtmlDocumentLoader),
	],
)
def test_factory_selects_correct_loader(tmp_path, extension, expected_loader_type):
	file_path = tmp_path / f"file{extension}"
	file_path.write_text("content", encoding="utf-8")
	loader = DocumentLoaderFactory.get_loader(file_path)
	assert isinstance(loader, expected_loader_type)

def test_factory_is_case_insensitive_on_extension(tmp_path):
	file_path = tmp_path / "FILE.TXT"
	file_path.write_text("content", encoding="utf-8")
	loader = DocumentLoaderFactory.get_loader(file_path)
	assert isinstance(loader, TxtDocumentLoader)

def test_factory_raises_for_unsupported_extension(tmp_path):
	file_path = tmp_path / "file.xyz"
	file_path.write_text("content", encoding="utf-8")
	with pytest.raises(UnsupportedFormatError):
		DocumentLoaderFactory.get_loader(file_path)


def test_factory_raises_filenotfound_for_missing_file(tmp_path):
	missing_path = tmp_path / "missing.txt"
	with pytest.raises(FileNotFoundError):
		DocumentLoaderFactory.get_loader(missing_path)
