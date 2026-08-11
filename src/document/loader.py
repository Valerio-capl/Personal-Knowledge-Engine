import hashlib
from abc import ABC, abstractmethod
from pathlib import Path
from document.exceptions import (
	DocumentParsingError,
	EncodingDetectionError,
)
from dataclasses import dataclass
from charset_normalizer import from_path
import fitz
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from bs4 import BeautifulSoup

@dataclass(frozen=True)
class FileMetadata:
	filename: str
	filepath: str
	extension: str
	size_bytes: int
	content_hash: str
	page_number:int | None = None

@dataclass(frozen=True)
class LoadedDocument:
	document_type: str
	content: str
	metadata: FileMetadata

class DocumentLoader(ABC):
	def __init__(self, file_path: str | Path):
		self.file_path = Path(file_path).resolve()
		if not self.file_path.is_file():
			raise FileNotFoundError(self.file_path)
		self._content_hash = self._compute_content_hash()
		
	@property
	def content_hash(self) -> str:
		return self._content_hash

	@abstractmethod
	def load(self) -> list[LoadedDocument]:
		""" load document """
		pass

	def _compute_content_hash(self) -> str:
		hasher = hashlib.sha256()
		try:
			with open(self.file_path, "rb") as f:
				for block in iter(lambda: f.read(65536), b""):
					hasher.update(block)
		except OSError as e:
			raise DocumentParsingError(f"Unable to read file {self.file_path} for hashing: {e}") from e
		return hasher.hexdigest()

	def _get_metadata(self,page_number:int = None) -> FileMetadata:
		stat = self.file_path.stat()
		return FileMetadata(
			filename=self.file_path.name,
			filepath=str(self.file_path),
			extension=self.file_path.suffix,
			size_bytes=stat.st_size,
			content_hash=self._content_hash,
			page_number=page_number
		)
	def _read_text_auto_encoding(self) -> str:
		try:
			result = from_path(str(self.file_path)).best()
		except OSError as e:
			raise EncodingDetectionError(f"Unable to read file {self.file_path}: {e}") from e
 
		if result is None:
			raise EncodingDetectionError(f"Unable to detect file encoding {self.file_path}")
		return str(result)


class TxtDocumentLoader(DocumentLoader):
	def load(self) -> list[LoadedDocument]:
		content=self._read_text_auto_encoding()
		return [LoadedDocument(
			document_type="txt",
			content=content,
			metadata=self._get_metadata(page_number=1)
		)]

class PdfDocumentLoader(DocumentLoader):
	def load(self) -> list[LoadedDocument]:
		documents = []
		try:
			with fitz.open(str(self.file_path)) as file:
				for page_idx, page in enumerate(file, start=1):
					try:
						text = page.get_text()
					except Exception as e:
						raise DocumentParsingError(f"Error extracting text from page {page_idx} of {self.file_path}: {e}") from e
					if not text.strip():
						continue
					documents.append(
						LoadedDocument(
							document_type="pdf",
							content=text,
							metadata=self._get_metadata(page_number=page_idx)
							))
		except Exception as e:
			raise DocumentParsingError(f"Unable to open the file {self.file_path}: {e}") from e

		return documents
	
class DocxDocumentLoader(DocumentLoader):
	def load(self) -> list[LoadedDocument]:
		try:
			doc = Document(str(self.file_path))
		except Exception as e:
			raise DocumentParsingError(f"Unable to open the file {self.file_path}: {e}") from e
		
		try:
			content = self._extract_content(doc)
		except Exception as e:
			raise DocumentParsingError(f"Error extracting content from {self.file_path}: {e}") from e
		
		return [LoadedDocument(
			document_type="docx",
			content=content,
			metadata=self._get_metadata(page_number=1)
		)]
	
	def _extract_content(self, doc: Document) -> str:
		""" Extracts the content of the document following the actual order of the body """
		parts = []
		tables_found = 0
 
		for block in doc.element.body.iterchildren():
			if block.tag.endswith("}p"):
				text = self._paragraph_text_from_element(doc, block)
				if text.strip():
					parts.append(text)
			elif block.tag.endswith("}tbl"):
				table = self._table_from_element(doc, block)
				if table is not None:
					table_text = self._table_to_text(table)
					if table_text.strip():
						parts.append(table_text)
						tables_found += 1
		return "\n".join(parts)
	
	@staticmethod
	def _paragraph_text_from_element(doc: Document, element) -> str:
		return Paragraph(element, doc).text
	@staticmethod
	def _table_from_element(doc: Document, element) -> Table | None:
		try:
			return Table(element, doc)
		except Exception:
			return None
	@staticmethod
	def _table_to_text(table: Table) -> str:
		rows_text = []
		for row in table.rows:
			cells_text = [cell.text.strip() for cell in row.cells]
			rows_text.append(" | ".join(cells_text))
		return "\n".join(rows_text)

class HtmlDocumentLoader(DocumentLoader):
	def load(self) -> list[LoadedDocument]:
		raw_html = self._read_text_auto_encoding()
		try:
			soup = BeautifulSoup(raw_html, "html.parser")
			content = soup.get_text(separator="\n", strip=True)
		except Exception as e:
			raise DocumentParsingError(f"HTML parsing error {self.file_path}: {e}") from e
		
		return [LoadedDocument(
			document_type="html",
			content=content,
			metadata=self._get_metadata(page_number=1)
		)]
